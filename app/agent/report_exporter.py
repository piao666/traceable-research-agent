"""Export generated Markdown reports to common downloadable formats."""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any


ROOT = Path(__file__).resolve().parents[2]
REPORTS_ROOT = ROOT / "workspace" / "reports"

_FORMAT_EXTENSIONS = {
    "markdown": "md",
    "docx": "docx",
    "pdf": "pdf",
}
_FORMAT_MEDIA_TYPES = {
    "markdown": "text/markdown",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}


@dataclass(frozen=True)
class ReportExportResult:
    run_id: str
    format: str
    report_path: str

    def to_dict(self) -> dict[str, Any]:
        return {
            "run_id": self.run_id,
            "format": self.format,
            "report_path": self.report_path,
        }


def normalize_report_format(value: str | None) -> str:
    normalized = (value or "markdown").strip().lower()
    if normalized in {"md", "markdown"}:
        return "markdown"
    if normalized in _FORMAT_EXTENSIONS:
        return normalized
    raise ValueError(f"Unsupported report format: {value}")


def report_filename(run_id: str, report_format: str) -> str:
    normalized = normalize_report_format(report_format)
    return f"research_report_{_safe_run_id(run_id)}.{_FORMAT_EXTENSIONS[normalized]}"


def report_media_type(report_format: str) -> str:
    return _FORMAT_MEDIA_TYPES[normalize_report_format(report_format)]


def resolve_report_path(report_path: str, report_root: Path = REPORTS_ROOT) -> Path:
    target = (ROOT / report_path).resolve()
    root = report_root.resolve()
    if root != target and root not in target.parents:
        raise ValueError("Report path escaped workspace/reports")
    return target


def read_report_markdown(report_path: str) -> str:
    return resolve_report_path(report_path).read_text(encoding="utf-8")


def export_report(
    run_id: str,
    markdown: str,
    report_format: str = "markdown",
    report_root: Path = REPORTS_ROOT,
) -> ReportExportResult:
    normalized = normalize_report_format(report_format)
    report_root.mkdir(parents=True, exist_ok=True)
    target = _report_path(report_root, run_id, normalized)
    if normalized == "markdown":
        target.write_text(markdown, encoding="utf-8")
    elif normalized == "docx":
        _write_docx(markdown, target)
    else:
        _write_pdf(markdown, target)
    return ReportExportResult(
        run_id=run_id,
        format=normalized,
        report_path=target.relative_to(ROOT).as_posix(),
    )


def _safe_run_id(run_id: str) -> str:
    return re.sub(r"[^A-Za-z0-9_-]", "_", run_id)[:96] or "unknown"


def _report_path(report_root: Path, run_id: str, report_format: str) -> Path:
    target = (
        report_root / f"{_safe_run_id(run_id)}.{_FORMAT_EXTENSIONS[report_format]}"
    ).resolve()
    root = report_root.resolve()
    if root != target and root not in target.parents:
        raise ValueError("Report path escaped workspace/reports")
    return target


def _iter_blocks(markdown: str) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    paragraph: list[str] = []
    in_code = False
    code_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph
        if paragraph:
            blocks.append({"type": "paragraph", "text": " ".join(paragraph).strip()})
            paragraph = []

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if line.strip().startswith("```"):
            if in_code:
                blocks.append({"type": "code", "text": "\n".join(code_lines)})
                code_lines = []
                in_code = False
            else:
                flush_paragraph()
                in_code = True
                code_lines = []
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip():
            flush_paragraph()
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            flush_paragraph()
            blocks.append(
                {
                    "type": "heading",
                    "level": len(heading.group(1)),
                    "text": heading.group(2).strip(),
                }
            )
            continue
        ordered = re.match(r"^\s*\d+[.)]\s+(.+)$", line)
        unordered = re.match(r"^\s*[-*]\s+(.+)$", line)
        if ordered or unordered:
            flush_paragraph()
            blocks.append(
                {
                    "type": "list",
                    "ordered": bool(ordered),
                    "text": (ordered or unordered).group(1).strip(),
                }
            )
            continue
        paragraph.append(line.strip())

    if in_code:
        blocks.append({"type": "code", "text": "\n".join(code_lines)})
    flush_paragraph()
    return blocks


def _clean_inline_markdown(text: str) -> str:
    cleaned = re.sub(r"`([^`]+)`", r"\1", text)
    cleaned = re.sub(r"\*\*([^*]+)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"\*([^*]+)\*", r"\1", cleaned)
    cleaned = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", cleaned)
    return cleaned


def _write_docx(markdown: str, target: Path) -> None:
    from docx import Document

    document = Document()
    for block in _iter_blocks(markdown):
        text = _clean_inline_markdown(str(block.get("text") or ""))
        if block["type"] == "heading":
            document.add_heading(text, level=min(int(block.get("level") or 1), 4))
        elif block["type"] == "list":
            style = "List Number" if block.get("ordered") else "List Bullet"
            document.add_paragraph(text, style=style)
        elif block["type"] == "code":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(str(block.get("text") or ""))
            run.font.name = "Consolas"
        elif text:
            document.add_paragraph(text)
    document.save(target)


def _write_pdf(markdown: str, target: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import Paragraph, Preformatted, SimpleDocTemplate, Spacer

    doc = SimpleDocTemplate(str(target), pagesize=A4)
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    for style in styles.byName.values():
        style.fontName = "STSong-Light"
    story: list[Any] = []
    for block in _iter_blocks(markdown):
        text = _clean_inline_markdown(str(block.get("text") or ""))
        if block["type"] == "heading":
            level = min(int(block.get("level") or 1), 3)
            story.append(Paragraph(text, styles[f"Heading{level}"]))
            story.append(Spacer(1, 8))
        elif block["type"] == "list":
            marker = "1. " if block.get("ordered") else "- "
            story.append(Paragraph(marker + text, styles["BodyText"]))
            story.append(Spacer(1, 4))
        elif block["type"] == "code":
            story.append(Preformatted(str(block.get("text") or ""), styles["Code"]))
            story.append(Spacer(1, 8))
        elif text:
            story.append(Paragraph(text, styles["BodyText"]))
            story.append(Spacer(1, 6))
    doc.build(story)
